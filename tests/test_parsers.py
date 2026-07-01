"""Tests for the document parsing layer (TXT path; PDF/DOCX need optional deps)."""

import pytest

from app.ingestion.parsers import (
    UnsupportedFormatError,
    detect_format,
    parse,
    parse_docx,
    parse_pdf,
    parse_txt,
)


def test_detect_format():
    assert detect_format("contract.pdf") == "pdf"
    assert detect_format("contract.DOCX") == "docx"
    assert detect_format("contract.txt") == "txt"


def test_detect_format_unsupported():
    with pytest.raises(UnsupportedFormatError):
        detect_format("contract.rtf")


def test_parse_txt_roundtrip():
    text = "Hello, contract."
    assert parse_txt(text.encode("utf-8")) == text


def test_parse_dispatches_txt():
    fmt, text = parse("nda.txt", b"1. Confidentiality\nKeep it secret.")
    assert fmt == "txt"
    assert "Confidentiality" in text


def _build_minimal_pdf(text: str) -> bytes:
    """Build a minimal single-page PDF containing `text`, with a correct
    byte-exact xref table (offsets are computed from the actual bytes
    written, not hand-counted) so pypdf can parse it without any extra
    PDF-generation dependency."""
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
            b"/MediaBox [0 0 400 200] /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    stream = f"BT /F1 12 Tf 20 150 Td ({text}) Tj ET".encode("latin-1")
    objects.append(
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
    )

    header = b"%PDF-1.4\n"
    offsets = [0]
    body = bytearray(header)
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"

    xref_offset = len(body)
    n = len(objects) + 1
    xref = f"xref\n0 {n}\n0000000000 65535 f \n".encode()
    for off in offsets[1:]:
        xref += f"{off:010d} 00000 n \n".encode()
    trailer = f"trailer\n<< /Size {n} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode()

    return bytes(body) + xref + trailer


def test_parse_pdf_extracts_text():
    pdf_bytes = _build_minimal_pdf("Sample Confidentiality Clause Text")
    text = parse_pdf(pdf_bytes)
    assert "Sample Confidentiality Clause Text" in text


def test_parse_docx_extracts_text():
    import io

    from docx import Document

    document = Document()
    document.add_paragraph("MUTUAL NON-DISCLOSURE AGREEMENT")
    document.add_paragraph("1. Confidentiality")
    document.add_paragraph("Each party agrees to hold in strict confidence all proprietary information.")
    buffer = io.BytesIO()
    document.save(buffer)

    text = parse_docx(buffer.getvalue())
    assert "MUTUAL NON-DISCLOSURE AGREEMENT" in text
    assert "Each party agrees to hold in strict confidence all proprietary information." in text


def test_parse_dispatches_pdf():
    pdf_bytes = _build_minimal_pdf("Dispatch check text")
    fmt, text = parse("contract.pdf", pdf_bytes)
    assert fmt == "pdf"
    assert "Dispatch check text" in text


def test_parse_dispatches_docx():
    import io

    from docx import Document

    document = Document()
    document.add_paragraph("Dispatch check text")
    buffer = io.BytesIO()
    document.save(buffer)

    fmt, text = parse("contract.docx", buffer.getvalue())
    assert fmt == "docx"
    assert "Dispatch check text" in text
